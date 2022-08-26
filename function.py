from docx import Document


def criar_matriz_curso(curso_titulo, curso_descricao, curso_professor, topico_1, topico_2, topico_3, topico_4, topico_5,
                       data_1, data_2, data_3, data_4, data_5,
                       arquivo_nome):
    document = Document()

    document.add_heading(curso_titulo, 0)

    document.add_paragraph(curso_descricao)

    document.add_heading('Professor', level=1)

    document.add_paragraph(curso_professor)

    document.add_heading('Ementa', level=1)

    document.add_paragraph(topico_1, style='List Bullet')
    document.add_paragraph(topico_2, style='List Bullet')
    document.add_paragraph(topico_3, style='List Bullet')
    document.add_paragraph(topico_4, style='List Bullet')
    document.add_paragraph(topico_5, style='List Bullet')

    document.add_heading('Cronograma', level=1)

    records = (
        (data_1, topico_1),
        (data_2, topico_2),
        (data_3, topico_3),
        (data_4, topico_4),
        (data_5, topico_5),
    )

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Data'
    hdr_cells[1].text = 'Tópico'
    for data, topico in records:
        row_cells = table.add_row().cells
        row_cells[0].text = data
        row_cells[1].text = topico

    document.save(arquivo_nome + '.docx')
